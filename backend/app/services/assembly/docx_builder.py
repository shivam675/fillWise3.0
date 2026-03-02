"""
Assembly Engine: reconstructs a DOCX from approved section rewrites.

Rules:
  - Only sections with an APPROVED review are included.
  - If a review has edited_text, that takes precedence over the LLM output.
  - Sections without a rewrite (no job cover) use the original text.
  - Assembly metadata is stored in DOCX core properties (not visible in body).
  - The output DOCX is written to the export directory.
  - Any residual Markdown formatting from the LLM is stripped before rendering.
"""

from __future__ import annotations

import json
import uuid
from datetime import UTC, datetime
from pathlib import Path
import structlog
from docx import Document  # type: ignore[import-untyped]
from docx.shared import Pt  # type: ignore[import-untyped]
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config.settings import get_settings
from app.core.errors import ConflictError, ErrorCode, NotFoundError
from app.db.models.document import Section
from app.db.models.job import JobStatus, RewriteJob, SectionRewrite
from app.db.models.review import Review, ReviewStatus
from app.services.llm.prompt_engine import _strip_trailing_metadata, strip_markdown

_log = structlog.get_logger(__name__)


class AssemblyEngine:
    """
    Assembles approved rewritten sections into a final DOCX.

    Validates that all rewrites have been reviewed before proceeding.
    """

    def __init__(self, db: AsyncSession) -> None:
        self._db = db
        self._settings = get_settings()

    async def assemble(self, job_id: str, requester_username: str) -> Path:
        """
        Build a DOCX from all approved sections for a completed job.

        Returns:
            Path to the exported DOCX file.

        Raises:
            ConflictError: If any sections are not yet reviewed/approved.
        """
        log = _log.bind(job_id=job_id)

        job = await self._db.get(RewriteJob, job_id)
        if job is None:
            raise NotFoundError("RewriteJob", job_id)

        if job.status != JobStatus.COMPLETED:
            raise ConflictError(
                ErrorCode.ASSEMBLY_FAILED,
                f"Job is in status '{job.status}'; assembly requires COMPLETED.",
            )

        # Load all sections in order
        section_result = await self._db.execute(
            select(Section)
            .where(Section.document_id == job.document_id)
            .order_by(Section.sequence_no)
        )
        sections: list[Section] = list(section_result.scalars().all())

        # Load all rewrites for this job
        rewrite_result = await self._db.execute(
            select(SectionRewrite).where(SectionRewrite.job_id == job_id)
        )
        rewrites_by_section: dict[str, SectionRewrite] = {
            r.section_id: r for r in rewrite_result.scalars().all()
        }

        # Load all reviews
        rewrite_ids = [r.id for r in rewrites_by_section.values()]
        reviews_by_rewrite: dict[str, Review] = {}
        if rewrite_ids:
            review_result = await self._db.execute(
                select(Review).where(Review.rewrite_id.in_(rewrite_ids))
            )
            reviews_by_rewrite = {r.rewrite_id: r for r in review_result.scalars().all()}

        # Pre-flight: ensure no PENDING reviews
        pending = [
            s.id for s in sections
            if s.id in rewrites_by_section
            and rewrites_by_section[s.id].id not in reviews_by_rewrite
        ]
        if pending:
            raise ConflictError(
                ErrorCode.ASSEMBLY_PENDING_REVIEWS,
                f"{len(pending)} section(s) have not been reviewed yet.",
            )

        non_approved = [
            s.id for s in sections
            if s.id in rewrites_by_section
            and (rw := rewrites_by_section[s.id])
            and (rev := reviews_by_rewrite.get(rw.id))
            and rev.status not in (ReviewStatus.APPROVED, ReviewStatus.EDITED)
        ]
        if non_approved:
            raise ConflictError(
                ErrorCode.ASSEMBLY_PENDING_REVIEWS,
                f"{len(non_approved)} section(s) have not been approved.",
            )

        # Build DOCX
        doc = Document()
        self._set_default_styles(doc)

        for section in sections:
            text = self._resolve_text(section, rewrites_by_section, reviews_by_rewrite)

            if section.heading and section.heading == section.original_text:
                # It's a standalone heading — strip any residual markdown bold
                clean_heading = strip_markdown(text).strip()
                para = doc.add_paragraph(clean_heading)
                para.style = doc.styles["Heading 1"]
            else:
                # Strip any residual markdown, then add as plain paragraphs
                clean_text = strip_markdown(text)
                for line in clean_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)

        # Store assembly metadata as DOCX core properties (not visible in body)
        core_props = doc.core_properties
        core_props.comments = json.dumps(
            {
                "job_id": job_id,
                "assembled_by": requester_username,
                "assembled_at": datetime.now(UTC).isoformat(),
                "fillwise_version": "3.0.0",
            }
        )

        # Write to export dir
        output_path = self._settings.export_dir / f"{job_id}_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(output_path))

        # Persist the filename so the job record knows where its export lives
        job.export_filename = output_path.name
        await self._db.flush()

        log.info("assembly_complete", output_path=str(output_path))
        return output_path

    def _resolve_text(
        self,
        section: Section,
        rewrites: dict[str, SectionRewrite],
        reviews: dict[str, Review],
    ) -> str:
        """Determine the final text for a section, precedence: edit > rewrite > original."""
        rewrite = rewrites.get(section.id)
        if rewrite is None:
            return section.original_text

        review = reviews.get(rewrite.id)
        if review is None:
            return section.original_text

        if review.edited_text:
            text = review.edited_text
        elif rewrite.rewritten_text:
            text = rewrite.rewritten_text
        else:
            return section.original_text

        # Safety net: strip any trailing LLM metadata that slipped through
        text = _strip_trailing_metadata(text, {})
        # Strip any residual markdown formatting
        text = strip_markdown(text)
        return text

    def _set_default_styles(self, doc: Document) -> None:
        """Configure default font for the output document."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
