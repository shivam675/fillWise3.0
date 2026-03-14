"""
DocumentProcessor: orchestrates ingestion of uploaded files.

Responsible for:
  1. MIME type validation
  2. File hash computation
  3. Page count enforcement
  4. Text extraction (PDF or DOCX)
  5. Structure detection
  6. Section persistence
  7. Status transitions with error recovery
"""

from __future__ import annotations

import hashlib

import structlog
from sqlalchemy.ext.asyncio import AsyncSession

from app.config.settings import get_settings
from app.core.errors import (
    NotFoundError,
    ValidationError,
)
from app.db.models.document import Document, DocumentStatus, Section, SectionType
from app.services.ingestion.docx_extractor import ExtractedParagraph, extract_docx
from app.services.ingestion.pdf_extractor import extract_pdf
from app.services.ingestion.structure_detector import StructuredSection, detect_structure

_log = structlog.get_logger(__name__)


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _text_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


class DocumentProcessor:
    """
    Processes a single uploaded document.

    Instances are not reused; create a new instance per call.
    """

    def __init__(self, db: AsyncSession) -> None:
        self._db = db
        self._settings = get_settings()

    async def process(self, document_id: str) -> None:
        """
        Full ingestion pipeline for an already-persisted Document record.

        Transitions document status:
          pending → extracting → mapping → mapped
               or → failed (on any error, with message stored)
        """
        log = _log.bind(document_id=document_id)

        document = await self._db.get(Document, document_id)
        if document is None:
            raise NotFoundError("Document", document_id)

        try:
            await self._set_status(document, DocumentStatus.EXTRACTING)
            log.info("ingestion_started")

            file_path = self._settings.upload_dir / document.filename
            raw_data = file_path.read_bytes()

            # Extract text per file type
            if document.mime_type == "application/pdf":
                pages = extract_pdf(raw_data)
                page_count = len(pages)
                paragraphs = self._pages_to_paragraphs(pages)
            else:
                content = extract_docx(raw_data)
                
                # Generate page content count
                def _get_docx_pages(raw: bytes) -> int | None:
                    import zipfile
                    import io
                    import lxml.etree as ET
                    
                    namespaces = {
                        "ap": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
                    }
                    try:
                        with zipfile.ZipFile(io.BytesIO(raw)) as docx:
                            app_xml = docx.read("docProps/app.xml")
                            root = ET.fromstring(app_xml)
                            pages_node = root.find(".//ap:Pages", namespaces)
                            if pages_node is not None and pages_node.text and pages_node.text.isdigit():
                                return int(pages_node.text)
                    except Exception:
                        pass
                    return None

                def _extract_docx_comments(raw: bytes) -> list[dict]:
                    import zipfile
                    import io
                    import lxml.etree as ET
                    from docx import Document
                    
                    namespaces = {
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    }
                    try:
                        comments_dict = {}
                        with zipfile.ZipFile(io.BytesIO(raw)) as docx:
                            try:
                                comments_xml = docx.read('word/comments.xml')
                                root = ET.fromstring(comments_xml)
                                for comment in root.findall('.//w:comment', namespaces):
                                    c_id = comment.get(f"{{{namespaces['w']}}}id")
                                    author = comment.get(f"{{{namespaces['w']}}}author")
                                    text = "".join(comment.itertext()).strip()
                                    if text:
                                        comments_dict[c_id] = {'author': author, 'text': text}
                            except KeyError:
                                return []
                        
                        if not comments_dict:
                            return []
                        
                        doc = Document(io.BytesIO(raw))
                        extracted_comments = []
                        for p in doc.paragraphs:
                            c_refs = p._element.xpath('.//w:commentRangeStart/@w:id')
                            if c_refs:
                                for c_id in set(c_refs):
                                    c = comments_dict.get(str(c_id))
                                    if c:
                                        extracted_comments.append({
                                            'author': c['author'],
                                            'text': c['text'],
                                            'context': p.text
                                        })
                        return extracted_comments
                    except Exception:
                        return []
                    
                page_count = _get_docx_pages(raw_data)
                paragraphs = content.paragraphs
                
                # Extract comments -> Generate Ruleset
                comments = _extract_docx_comments(raw_data)
                if comments:
                    try:
                        import json
                        from app.db.models.ruleset import Ruleset
                        from app.services.rules.validator import compute_rules_hash
                        from sqlalchemy import select

                        rules_list = []
                        
                        try:
                            from app.services.llm.client import OllamaClient
                            
                            client = OllamaClient()
                            system_prompt = (
                                "You are a professional AI assistant that converts document review comments into a formal Ruleset JSON.\n"
                                "You will be given a list of comments and their context paragraphs.\n"
                                "Generate meaningful rules out of them and output ONLY a valid JSON array of objects without any markdown blocks around it.\n"
                                "Each object must have: 'id' (a unique string), 'name' (a short rule title), "
                                "and 'instruction' (the comprehensive instruction formatted as a clear mandate, at least 15 characters long)."
                            )
                            user_prompt_data = json.dumps(comments, indent=2)
                            response = await client.complete(system_prompt, f"Comments:\n{user_prompt_data}")
                            
                            llm_out = response.content.strip()
                            if llm_out.startswith("```json"):
                                llm_out = llm_out[7:-3].strip()
                            elif llm_out.startswith("```"):
                                llm_out = llm_out[3:-3].strip()
                                
                            parsed_rules = json.loads(llm_out)
                            for r in parsed_rules:
                                if isinstance(r, dict) and "instruction" in r and len(str(r["instruction"])) >= 10:
                                    rules_list.append({
                                        "id": str(r.get("id", f"ai-rule-{len(rules_list)+1}")),
                                        "name": str(r.get("name", "AI Extracted Rule")),
                                        "instruction": str(r["instruction"])
                                    })
                        except Exception as e:
                            _log.warning("ollama_comment_ruleset_failed", error=str(e))
                            # Fallback if Ollama fails or parsing fails
                            rules_list = []
                            for idx, c in enumerate(comments, 1):
                                rule_id = f"comment-rule-{idx}"
                                author = c.get('author') or "Reviewer"
                                instruction = (
                                    f"Apply the following comment: '{c['text']}'. "
                                    f"Context reference: '{c['context'][:100]}...'"
                                )
                                rules_list.append({
                                    "id": rule_id,
                                    "name": f"Comment from {author} #{idx}",
                                    "instruction": instruction
                                })

                        if rules_list:
                            # Use document ID to ensure unique but deterministic name
                            rs_name = f"Extracted Comments: {document.filename}"
                            rs_version = "1.0.0"
                            rs_dict = {
                                "name": rs_name,
                                "version": rs_version,
                                "rules": rules_list
                            }
                            
                            existing = await self._db.execute(
                                select(Ruleset).where(Ruleset.name == rs_name)
                            )
                            if not existing.scalars().first():
                                rs = Ruleset(
                                    name=rs_name,
                                    description=f"Auto-generated ruleset from comments embedded inside {document.filename}",
                                    version=rs_version,
                                    schema_version="1.0",
                                    content_hash=compute_rules_hash(rs_dict),
                                    is_active=True,  # Activate it immediately
                                    rules_json=json.dumps(rules_list),
                                    created_by=document.created_by,
                                )
                                self._db.add(rs)
                    except Exception as e:
                        _log.warning("failed_to_extract_comments_ruleset", error=str(e))

            # Enforce page limit
            # PDF: exact page count; DOCX: estimate ~3000 chars per page (legal docs)
            if page_count is not None:
                effective_pages = page_count
            else:
                total_chars = sum(len(p.text) for p in paragraphs)
                effective_pages = max(1, total_chars // 3000)
            if effective_pages > self._settings.max_document_pages:
                raise ValidationError(
                    f"Document has {effective_pages} pages which exceeds the limit of "
                    f"{self._settings.max_document_pages}",
                    detail={"pages": effective_pages, "limit": self._settings.max_document_pages},
                )

            document.page_count = page_count

            await self._set_status(document, DocumentStatus.MAPPING)
            sections = detect_structure(paragraphs)
            await self._persist_sections(document, sections)

            await self._set_status(document, DocumentStatus.MAPPED)
            log.info("ingestion_complete", sections=len(sections))

        except Exception as exc:
            log.error("ingestion_failed", error=str(exc))
            document.status = DocumentStatus.FAILED
            document.error_message = str(exc)[:1000]
            await self._db.flush()
            raise

    def _pages_to_paragraphs(
        self, pages: list[tuple[int, str]]
    ) -> list[ExtractedParagraph]:
        """Convert PDF page tuples into ExtractedParagraph objects for detect_structure."""
        paragraphs: list[ExtractedParagraph] = []
        idx = 0
        for _page_no, text in pages:
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    paragraphs.append(
                        ExtractedParagraph(
                            text=line,
                            paragraph_index=idx,
                            style_name="Normal",
                            is_bold=False,
                            is_italic=False,
                        )
                    )
                    idx += 1
        return paragraphs

    async def _set_status(self, document: Document, status: DocumentStatus) -> None:
        document.status = status
        await self._db.flush()

    async def _persist_sections(
        self, document: Document, structured: list[StructuredSection]
    ) -> None:
        """
        Persist detected sections as Section records.

        Heading sections are used to assign parent_id to subsequent
        non-heading sections, building a simple two-level hierarchy.
        """
        current_heading_id: str | None = None

        for seq_no, s in enumerate(structured, start=1):
            section = Section(
                document_id=document.id,
                sequence_no=seq_no,
                section_type=s.section_type,
                heading=s.heading,
                original_text=s.text,
                content_hash=_text_hash(s.text),
                depth=s.depth,
                char_count=len(s.text),
            )

            if s.section_type == SectionType.HEADING:
                section.parent_id = None
                current_heading_id = None  # will be set after flush
            else:
                section.parent_id = current_heading_id

            self._db.add(section)
            await self._db.flush()

            if s.section_type == SectionType.HEADING:
                current_heading_id = section.id
